import os
import docx
from translatepy import Translator

os.chdir('C:/Users/Ab/VsCode Projects/Projects/Table Translator PSG')   #changes the current working directory of the system to the desired directory.
doc = docx.Document()
doc.save('Table Document.docx')   #redundant save statement that is used to create a file in case it didnt exist before.
translator_object = Translator()

title = doc.add_heading('The Table')
table = doc.add_table(rows = 1, cols = 3)   #the row containing column names(first row) is created with 3 columns during the creation of the table.
header_row = table.rows[0].cells   #stores the individual cells of the first row in a variable.
header_row[0].text = 'S.No'
header_row[1].text = 'Water'
header_row[2].text = 'Fire'

data = (
    (1, 'yes', 'no'),
    (2, 'no', 'yes'),
)
#stores sno, choice1, choice2 from the nested tuples in the 'data' tuple to the respective cells of the newly created row.
for sno, choice1, choice2 in data:
    new_row = table.add_row()
    new_cells = new_row.cells
    new_cells[0].text = str(sno)
    new_cells[1].text = choice1
    new_cells[2].text = choice2
table.style = 'Colorful List'   #stylizes the table to the 'Colorful List' type.
for n,row in enumerate(table.rows):
    for m,cell in enumerate(row.cells):
        translation = translator_object.translate(cell.text,'ta')   #translates the text of the individual cells to the desired language.
        table.cell(n,m).text = str(translation)   #str type casting is necessary since the text function automatically assumes the TranslationResult type
doc.save('Table Document.docx')
